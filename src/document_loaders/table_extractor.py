"""
RAG Pipeline — Table Extractor

Extracts tables from PDFs and DOCX files, converting them to
Markdown format for better LLM comprehension.
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from src.config import ExtractionConfig


def table_to_markdown(table: list[list[str | None]]) -> str:
    """Convert a 2D table (list of rows) to Markdown format.

    Args:
        table: List of rows, where each row is a list of cell values.

    Returns:
        Markdown-formatted table string.
    """
    if not table or not table[0]:
        return ""

    # Clean cells: replace None with empty string, strip whitespace
    cleaned = []
    for row in table:
        cleaned.append([str(cell).strip() if cell else "" for cell in row])

    # Normalize column count
    max_cols = max(len(row) for row in cleaned)
    for row in cleaned:
        while len(row) < max_cols:
            row.append("")

    # Build Markdown table
    header = cleaned[0]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in cleaned[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def extract_tables_from_pdf(
    pdf_path: str | Path,
    config: ExtractionConfig,
) -> list[Document]:
    """Extract tables from a PDF file using pdfplumber.

    Args:
        pdf_path: Path to the PDF file.
        config: Extraction configuration.

    Returns:
        List of Document objects, one per table found.
    """
    import pdfplumber

    pdf_path = Path(pdf_path)
    documents: list[Document] = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()

                for table_idx, table in enumerate(tables or []):
                    if not table or len(table) < config.min_table_rows:
                        continue

                    md = table_to_markdown(table)
                    if not md.strip():
                        continue

                    doc = Document(
                        page_content=md,
                        metadata={
                            "source": str(pdf_path),
                            "type": "table",
                            "page": page_num,
                            "table_index": table_idx,
                        },
                    )
                    documents.append(doc)

    except Exception as e:
        print(f"  [ERROR] Error extracting tables from {pdf_path.name}: {e}")

    return documents


def extract_tables_from_docx(
    docx_path: str | Path,
    config: ExtractionConfig,
) -> list[Document]:
    """Extract tables from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file.
        config: Extraction configuration.

    Returns:
        List of Document objects, one per table found.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("  [WARN]  python-docx not installed. Skipping DOCX table extraction.")
        return []

    docx_path = Path(docx_path)
    documents: list[Document] = []

    try:
        doc = DocxDocument(str(docx_path))
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])

            if len(rows) < config.min_table_rows:
                continue

            md = table_to_markdown(rows)
            if not md.strip():
                continue

            doc = Document(
                page_content=md,
                metadata={
                    "source": str(docx_path),
                    "type": "table",
                    "page": 0,
                    "table_index": table_idx,
                },
            )
            documents.append(doc)

    except Exception as e:
        print(f"  [ERROR] Error extracting tables from {docx_path.name}: {e}")

    return documents
